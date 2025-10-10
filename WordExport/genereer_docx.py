import os
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import locale

# ======================
# Opmaak-helpers (alleen layout, geen logica)
# ======================

def _tighten_paragraph(paragraph):
    """Alinea-afstand overal 0; enkelvoudige regelafstand."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

def _enforce_table_style(table):
    """Zet alle tekst in de tabel op 10pt en alinea-afstand 0."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _tighten_paragraph(para)
                for run in para.runs:
                    run.font.size = Pt(10)

def _add_paragraph(doc, text=""):
    """Voeg paragraaf toe met strakke spacing (0)."""
    p = doc.add_paragraph(text if text is not None else "")
    _tighten_paragraph(p)
    return p

def _as_str(x, default=""):
    """Zachte cast naar string (voorkomt typefouten)."""
    if x is None:
        return default
    try:
        s = str(x)
    except Exception:
        return default
    return s

def _as_list(x):
    """Zachte cast naar list (None -> [])."""
    if x is None:
        return []
    if isinstance(x, list):
        return x
    # strings niet uit elkaar trekken; behandel als enkel item
    return [x]

# ======================
# Bestaande functies (ongewijzigd qua gedrag)
# ======================

def maak_in_klapbare_heading(paragraph, text):
    run = paragraph.add_run(text)
    rPr = run._r.get_or_add_rPr()
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Heading3')
    rPr.append(rStyle)

def collapse_heading(paragraph, collapsed=True):
    """
    Zet 'Collapsed by default' op een heading-paragraaf (Word 2013+).
    paragraph: python-docx Paragraph die al een Heading-stijl heeft.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # Verwijder bestaande w15:collapsed nodes (als je herhaald aanroept)
    for child in list(pPr):
        if child.tag == '{http://schemas.microsoft.com/office/word/2012/wordml}collapsed':
            pPr.remove(child)

    if collapsed:
        node = parse_xml(
            r'<w15:collapsed xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>'
        )
        pPr.append(node)

# ======================
# Hoofdfunctie (zelfde signatuur/functionaliteit)
# ======================

def genereer_word_document(patiënten_data, afdeling, output_path: str | None = None):
    """
    Schrijft naar `output_path` als opgegeven, anders naar Output/MedicatieReview_<afdeling>.docx.
    Functionaliteit blijft exact gelijk aan je originele script.
    Alleen opmaak (0 spacing overal, 10pt in tabellen) en robuustheid (None/typen) zijn verbeterd.
    + Harde pagina-einde (Ctrl+Enter) tussen patiënten.
    """
    doc = Document()

        # Stel Nederlandse locale in (werkt op de meeste systemen)
    try:
        locale.setlocale(locale.LC_TIME, 'nl_NL.UTF-8')
    except locale.Error:
        # fallback als locale niet beschikbaar is (Windows of sommige servers)
        pass

    # Globale stijl: spacing 0; (laat font-size Normal verder met default)
    try:
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        # font-size Normal niet forceren; alleen spacing
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
    except Exception:
        pass  # als stijl niet beschikbaar, gaan we door

    
    # Standaard documenttaal instellen op Nederlands
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        styles = doc.styles
        for style in styles:
            if style.type == 1:  # paragraph style
                lang = OxmlElement('w:lang')
                lang.set(qn('w:val'), 'nl-NL')  # Nederlands (Nederland)
                style.element.rPr.append(lang)
    except Exception as e:
        print(f"Waarschuwing: kon taal niet instellen ({e})")

    # Marges en logo
    try:
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        _tighten_paragraph(header_para)
        header_para.alignment = 2  # Rechts uitlijnen
        run = header_para.add_run()
        logo_path = os.path.join("Data", "logo_apotheek_rgb.jpg")
        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Cm(4))
        else:
            print(f"Waarschuwing: Logo niet gevonden op {logo_path}")
    except Exception as e:
        print(f"Waarschuwing: Fout bij toevoegen van logo: {str(e)}")

    # Hoofdtitel
    heading1 = doc.add_heading(f"Medicatiebeoordeling - Afdeling {_as_str(afdeling)}", level=1)
    try:
        heading1.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    except Exception:
        pass
    _tighten_paragraph(heading1)

    vandaag = datetime.today().strftime("%d-%m-%Y")

    # Maak een stabiele lijst en gebruik enumerate voor page breaks tussen patiënten
    patients_list = _as_list(patiënten_data)

    for idx, patiënt in enumerate(patients_list):
        # Zorg dat we een dict hebben
        patiënt = patiënt if isinstance(patiënt, dict) else {}

        naam = _as_str(patiënt.get("naam", "Onbekend"), "Onbekend")
        heading2 = doc.add_heading(naam, level=2)
        try:
            heading2.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        _tighten_paragraph(heading2)

        # Arts, apotheker, datum, eGFR
        para = _add_paragraph(doc)
        for label, value in [("Arts:", ""), ("Apotheker:", ""), ("Datum:", vandaag), ("eGFR:", "")]:
            run = para.add_run(f"{label} ")
            run.bold = True
            para.add_run(_as_str(value) + "\n")

        # ====== STOPP criteria ======
        h_stopp = doc.add_heading("Mogelijke STOPP-criteria:", level=3)
        try:
            h_stopp.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        collapse_heading(h_stopp, True)
        _tighten_paragraph(h_stopp)

        stopp_items = _as_list(patiënt.get("stopp"))
        if stopp_items:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Criteriumcode", "Categorie", "Beschrijving", "Argument", "Getriggerd door"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(_as_str(text))
                run.bold = True

            for item in stopp_items:
                item = item if isinstance(item, dict) else {}
                row_cells = table.add_row().cells
                row_cells[0].text = _as_str(item.get('id', ''))
                row_cells[1].text = _as_str(item.get('category', ''))
                row_cells[2].text = _as_str(item.get('description', ''))
                row_cells[3].text = _as_str(item.get('argument', ''))
                row_cells[4].text = _as_str(item.get('triggering_medicines', ''))
            _enforce_table_style(table)
        else:
            _add_paragraph(doc, "Geen STOPP-criteria getriggerd.")

        # ====== Dubbelmedicatie ======
        h_dubbel = doc.add_heading("Mogelijke dubbelmedicatie:", level=3)
        try:
            h_dubbel.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        collapse_heading(h_dubbel, True)
        _tighten_paragraph(h_dubbel)

        dubbel_items = _as_list(patiënt.get("dubbelmedicatie"))
        if dubbel_items:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Geneesmiddelgroep (ATC)", "Geneesmiddelen"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(_as_str(text)); run.bold = True

            for item in dubbel_items:
                item = item if isinstance(item, dict) else {}
                row_cells = table.add_row().cells
                row_cells[0].text = _as_str(item.get('groep', "Onbekend"), "Onbekend")

                middelen = item.get('middelen')
                if isinstance(middelen, list):
                    geneesmiddelen_text = ", ".join(_as_str(m) for m in middelen if m not in (None, ""))
                elif isinstance(middelen, str):
                    geneesmiddelen_text = middelen
                else:
                    geneesmiddelen_text = _as_str(middelen) if middelen else "Geen middelen"
                row_cells[1].text = geneesmiddelen_text if geneesmiddelen_text else "Geen middelen"
            _enforce_table_style(table)
        else:
            _add_paragraph(doc, "Geen dubbelmedicatie gevonden.")

        # ====== ACB-score ======
        h_acb = doc.add_heading("Anticholinerge belastingscore (ACB-score):", level=3)
        try:
            h_acb.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        collapse_heading(h_acb, True)
        _tighten_paragraph(h_acb)

        acb_data = patiënt.get("acb")
        # Verwacht: (score, interpretatie, middelen_met_bijdrage)
        is_tuple = isinstance(acb_data, (list, tuple)) and len(acb_data) >= 2
        if is_tuple:
            score = _as_str(acb_data[0], "n.v.t.")
            interpretatie = _as_str(acb_data[1], "")
            middelen_met_bijdrage = _as_list(acb_data[2] if len(acb_data) > 2 else [])

            para = _add_paragraph(doc)
            run = para.add_run("Totale score: "); run.bold = True
            para.add_run(f"{score}" + (f" ({interpretatie})" if interpretatie else ""))

            if middelen_met_bijdrage:
                para = _add_paragraph(doc)
                run = para.add_run("Bijdragende geneesmiddelen:"); run.bold = True

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ["Geneesmiddel", "ACB-Score"]
                for i, text in enumerate(headers):
                    run = hdr_cells[i].paragraphs[0].add_run(_as_str(text)); run.bold = True

                for middel_info in middelen_met_bijdrage:
                    middel_info = middel_info if isinstance(middel_info, dict) else {}
                    row_cells = table.add_row().cells
                    row_cells[0].text = _as_str(middel_info.get('middel', ''))
                    row_cells[1].text = _as_str(middel_info.get('score', ''))
                _enforce_table_style(table)
            else:
                _add_paragraph(doc, "Geen bijdragende middelen.")
        else:
            para = _add_paragraph(doc)
            run = para.add_run("Totale score: "); run.bold = True
            para.add_run("n.v.t.")

        # ====== Vallen? (placeholder sectie) ======
        h_vallen = doc.add_heading("Vallen?", level=3)
        try:
            h_vallen.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        _tighten_paragraph(h_vallen)

        # ====== Malen? (extra placeholder sectie) ======
        h_malen = doc.add_heading("Malen?\n", level=3)
        try:
            h_malen.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        _tighten_paragraph(h_malen)

        # ====== Medicatieoverzicht per groep ======
        h_med = doc.add_heading("Medicatieoverzicht:", level=3)
        try:
            h_med.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        except Exception:
            pass
        _tighten_paragraph(h_med)

        groepen_dict = defaultdict(list)
        for gm in _as_list(patiënt.get("geneesmiddelen")):
            gm = gm if isinstance(gm, dict) else {}
            jansen_omschrijving = gm.get("ATC3_jansen")
            key = _as_str(jansen_omschrijving, "").strip() or "Overig"
            groepen_dict[key].append(gm)

        # Sorteer stabiel en veilig (keys kunnen non-str zijn)
        def _sort_key(k):
            s = _as_str(k).lower()
            return (s == "overig", s)
        gesorteerde_keys = sorted(groepen_dict.keys(), key=_sort_key)

        for jansen_omschrijving in gesorteerde_keys:
            middelen = groepen_dict.get(jansen_omschrijving, [])
            subheading = doc.add_heading(_as_str(jansen_omschrijving, "Overig"), level=4)
            try:
                subheading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
            except Exception:
                pass
            _tighten_paragraph(subheading)

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Geneesmiddel", "Geneesmiddelgroep (ATC)", "Gebruik", "Opmerking in Medimo"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(_as_str(text)); run.bold = True

            for gm in _as_list(middelen):
                gm = gm if isinstance(gm, dict) else {}
                row_cells = table.add_row().cells
                row_cells[0].text = _as_str(gm.get("clean", "-"), "-")
                atc5_omsch = gm.get("ATC5_omschrijving")
                groep_str = _as_str(atc5_omsch, "").strip() or "-"
                row_cells[1].text = groep_str
                row_cells[2].text = _as_str(gm.get("gebruik", "-"), "-")
                row_cells[3].text = _as_str(gm.get("opmerking", ""), "")
            _enforce_table_style(table)

            # Labels zoals in je laatste code-aanpassing
            for label in ["Eerder besproken:", "Opmerkingen:\n"]:
                para = _add_paragraph(doc)
                run = para.add_run(label)
                run.bold = True

        # ====== Harde pagina-einde tussen patiënten (Ctrl+Enter) ======
        if idx < len(patients_list) - 1:
            doc.add_page_break()


    # Outputpad kiezen
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
    else:
        os.makedirs("Output", exist_ok=True)
        doc_path = os.path.join("Output", f"MedicatieReview_{_as_str(afdeling)}_{datetime.today().strftime('%b%Y')}.docx")
        doc.save(doc_path)
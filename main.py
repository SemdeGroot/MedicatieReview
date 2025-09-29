import os
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# ===== Medimo Parser =====
from Parsers import parse_medimo

# ===== Externe analyses  =====
from START_STOP.check_start_stop import check_stopp_criteria
from Anticholinerge_Score.check_acb import bereken_acb_score
from Dubbelmedicatie.check_dubbelmedicatie import check_dubbelmedicatie


def maak_in_klapbare_heading(paragraph, text):
    run = paragraph.add_run(text)
    rPr = run._r.get_or_add_rPr()
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Heading3')
    rPr.append(rStyle)

def collapse_heading(paragraph, collapsed=True):
    """
    Zet 'Collapsed by default' op een heading-paragraaf (Word 2013+).
    paragraph: python-docx Paragraph die al een Heading-stijl heeft.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # Verwijder bestaande w15:collapsed nodes (als je herhaald aanroept)
    for child in list(pPr):
        if child.tag == '{http://schemas.microsoft.com/office/word/2012/wordml}collapsed':
            pPr.remove(child)

    if collapsed:
        node = parse_xml(
            r'<w15:collapsed xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>'
        )
        pPr.append(node)

def genereer_word_document(patiënten_data, afdeling):
    doc = Document()

    # Marges en logo
    try:
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = 2  # Rechts uitlijnen
        run = header_para.add_run()
        logo_path = os.path.join("Data", "logo_apotheek_rgb.jpg")
        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Cm(4))
        else:
            print(f"Waarschuwing: Logo niet gevonden op {logo_path}")
    except Exception as e:
        print(f"Waarschuwing: Fout bij toevoegen van logo: {str(e)}")

    # Hoofdtitel
    doc.add_heading(f"Medicatiebeoordeling - Afdeling {afdeling}", level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)

    vandaag = datetime.today().strftime("%d-%m-%Y")

    for patiënt in patiënten_data:
        heading = doc.add_heading(f"{patiënt['naam']}", level=2)
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)

        # Arts, apotheker, datum, eGFR
        para = doc.add_paragraph()
        for label, value in [("Arts:", ""), ("Apotheker:", ""), ("Datum:", vandaag), ("eGFR:", "")]:
            run = para.add_run(f"{label} ")
            run.bold = True
            para.add_run(f"{value}\n")

        # ====== STOPP criteria (tijdelijk uit) ======
        heading = doc.add_heading("Mogelijke STOPP-criteria:", level=3)
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        collapse_heading(heading, True)

        if patiënt.get("stopp"):
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Criteriumcode", "Categorie", "Beschrijving", "Argument", "Getriggerd door"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(text)
                run.bold = True

            for item in patiënt["stopp"]:
                row_cells = table.add_row().cells
                row_cells[0].text = item['id']
                row_cells[1].text = item['category']
                row_cells[2].text = item['description']
                row_cells[3].text = item['argument']
                row_cells[4].text = item['triggering_medicines']
        else:
            doc.add_paragraph("Geen STOPP-criteria getriggerd.")

        # ====== Dubbelmedicatie ======
        heading = doc.add_heading("Mogelijke dubbelmedicatie:", level=3)
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        collapse_heading(heading, True)

        if patiënt.get("dubbelmedicatie"):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Geneesmiddelgroep (ATC)", "Geneesmiddelen"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(text)
                run.bold = True

            for item in patiënt["dubbelmedicatie"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('groep')) if item.get('groep') else "Onbekend"
                middelen = item.get('middelen', [])
                if isinstance(middelen, list):
                    geneesmiddelen_text = ", ".join(str(m) for m in middelen if m)
                elif isinstance(middelen, str):
                    geneesmiddelen_text = middelen
                else:
                    geneesmiddelen_text = str(middelen) if middelen else "Geen middelen"
                row_cells[1].text = geneesmiddelen_text if geneesmiddelen_text else "Geen middelen"
        else:
            doc.add_paragraph("Geen dubbelmedicatie gevonden.")

        # ====== ACB-score  ======
        heading = doc.add_heading("Anticholinerge belastingscore (ACB-score):", level=3)
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        collapse_heading(heading, True)

        if patiënt.get("acb"):
            score, interpretatie, middelen_met_bijdrage = patiënt["acb"]
            para = doc.add_paragraph()
            run = para.add_run("Totale score: ")
            run.bold = True
            para.add_run(f"{score} ({interpretatie})")

            if middelen_met_bijdrage:
                para = doc.add_paragraph()
                run = para.add_run("Bijdragende geneesmiddelen:")
                run.bold = True

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ["Geneesmiddel", "ACB-Score"]
                for i, text in enumerate(headers):
                    run = hdr_cells[i].paragraphs[0].add_run(text)
                    run.bold = True

                for middel_info in middelen_met_bijdrage:
                    row_cells = table.add_row().cells
                    row_cells[0].text = middel_info['middel']
                    row_cells[1].text = str(middel_info['score'])
            else:
                doc.add_paragraph("Geen bijdragende middelen.")
        else:
            # Geen ACB berekend → toon lege sectie
            para = doc.add_paragraph()
            run = para.add_run("Totale score: ")
            run.bold = True
            para.add_run("n.v.t.")

        # ====== Medicatieoverzicht per groep (ATC/Jansen) ======
        heading = doc.add_heading("Medicatieoverzicht:", level=3)
        heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)

        # Groeperen op Jansen-omschrijving (uit ATC_groepen.db), fallback Overig
        groepen_dict = defaultdict(list)
        for gm in patiënt["geneesmiddelen"]:
            jansen_omschrijving = gm.get("ATC3_jansen") or "Overig"
            groepen_dict[jansen_omschrijving].append(gm)

        gesorteerde_keys = sorted(groepen_dict.keys(), key=lambda k: (k == "Overig", k.lower()))

        for jansen_omschrijving in gesorteerde_keys:
            middelen = groepen_dict[jansen_omschrijving]
            heading = doc.add_heading(f"{jansen_omschrijving}", level=4)
            heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x80)

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Geneesmiddel", "Geneesmiddelgroep (ATC)", "Gebruik", "Opmerking in Medimo"]
            for i, text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(text)
                run.bold = True

            for gm in middelen:
                row_cells = table.add_row().cells
                row_cells[0].text = gm.get("clean", "-")
                # Toon ATC5 + omschrijving samen als 'groep'
                atc5_omsch = gm.get("ATC5_omschrijving")
                groep_str = (f"{atc5_omsch}" if atc5_omsch else "-")
                row_cells[1].text = groep_str
                row_cells[2].text = gm.get("gebruik", "-")
                row_cells[3].text = gm.get("opmerking", "")

            # “Eerder besproken / Opmerking …” blok
            para = doc.add_paragraph()
            run = para.add_run("Eerder besproken:")
            run.bold = True
            para = doc.add_paragraph()
            run = para.add_run("Opmerking apotheker:")
            run.bold = True
            para = doc.add_paragraph()
            run = para.add_run("Opmerking arts:\n")
            run.bold = True

    os.makedirs("Output", exist_ok=True)
    doc_path = f"Output/MedicatieReview_{afdeling}.docx"
    doc.save(doc_path)

def main():
    # Nieuwe parser: retourneert (resultaat, afdeling)
    data, afdeling = parse_medimo.run_parser()

    patiënten_data = []
    for patiënt in data:
        naam = patiënt["patiënt"]

        middelen_clean = []
        medicatielijst = []  # voor latere analyses (STOPP/ACB/dubbel)
        for gm in patiënt["geneesmiddelen"]:
            # gm bevat al: NMNR, HPKode, SPKode, ATC, ATC3, ATC3_key, ATC3_omschrijving, ATC3_jansen
            # Voeg niets meer uit geneesmiddelen.db toe.
            # Voor een herkenbare naam gebruiken we de Medimo-clean naam.
            herkenbare_naam = gm.get("clean") or "Onbekend middel"

            medicatielijst.append(herkenbare_naam)
            middelen_clean.append(gm)

        # Externe analyses (tijdelijk enkele uitgeschakeld)
        leeftijd = 75  # Of dynamisch uitlezen indien beschikbaar
        stopp = check_stopp_criteria(middelen_clean, leeftijd)
        acb_score, interpretatie, middelen_met_bijdrage = bereken_acb_score(middelen_clean)
        acb = (acb_score, interpretatie, middelen_met_bijdrage)
        dubbel = check_dubbelmedicatie(middelen_clean)

        patiënten_data.append({
            "naam": naam,
            "geneesmiddelen": middelen_clean,
            "stopp": stopp,  
            "acb": acb,  # tijdelijk leeg
            "dubbelmedicatie": dubbel 
        })

    genereer_word_document(patiënten_data, afdeling)


if __name__ == "__main__":
    print("Bezig met analyseren...")
    main()
    print(f"Klaar! Word-document in Output/")